# -*- coding: utf-8 -*-
"""
请示函生成器
基于 Word 模板生成内部请示文档
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
import shutil
import re
import os
import sys
import subprocess

def convert_to_pdf(docx_path):
    """将 Word 文档转换为 PDF"""
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    try:
        # 尝试使用 docx2pdf
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return pdf_path
    except ImportError:
        pass
    
    try:
        # 尝试使用 unoconv (Linux/Mac)
        subprocess.run(['unoconv', '-f', 'pdf', '-o', pdf_path, docx_path], check=True)
        return pdf_path
    except (FileNotFoundError, subprocess.CalledProcessError):
        pass
    
    # 尝试使用 LibreOffice (Linux/Mac/Windows)
    try:
        # Windows: 使用 LibreOffice 的 headless 模式
        result = subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf', '--outdir', 
            os.path.dirname(docx_path), docx_path
        ], capture_output=True, timeout=30)
        if result.returncode == 0 and os.path.exists(pdf_path):
            return pdf_path
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    
    raise RuntimeError("无法转换 PDF，请安装 docx2pdf 或 LibreOffice")

def set_font(run, font_name='仿宋_GB2312', font_size=None):
    """设置字体"""
    if font_size:
        run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def convert_to_chinese_date(date_str):
    """将日期转换为中文格式"""
    if date_str == "今天":
        date_str = datetime.now().strftime("%Y年%m月%d日")
    
    numbers = re.findall(r'\d+', date_str)
    if len(numbers) >= 3:
        year, month, day = numbers[:3]
        chinese_nums = {'0': '〇', '1': '一', '2': '二', '3': '三', '4': '四', 
                       '5': '五', '6': '六', '7': '七', '8': '八', '9': '九'}
        year_chinese = ''.join(chinese_nums.get(c, c) for c in year)
        month_chinese = ''.join(chinese_nums.get(c, c) for c in month)
        day_chinese = ''.join(chinese_nums.get(c, c) for c in day)
        return f"{year_chinese}年{month_chinese}月{day_chinese}日"
    return date_str

def replace_placeholder(paragraph, placeholder, new_text, font_size=15, font_name='仿宋_GB2312'):
    """替换段落中的占位符"""
    full_text = paragraph.text
    if placeholder in full_text:
        # 清除原段落
        paragraph.clear()
        # 替换占位符
        new_text = full_text.replace(placeholder, new_text)
        run = paragraph.add_run(new_text)
        set_font(run, font_name, font_size)
        return True
    return False

def replace_table_placeholder(table, placeholder, new_text):
    """替换表格中的占位符"""
    for row in table.rows:
        for cell in row.cells:
            if placeholder in cell.text:
                cell.text = cell.text.replace(placeholder, new_text)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font(run, '仿宋_GB2312', 15)
                return True
    return False

def generate_request_letter(
    title,
    content,
    letter_number,
    date="今天",
    agent="游洋洋",
    template_path=None,
    output_path=None
):
    """
    使用模板生成请示函
    
    参数:
        title: 事项标题
        content: 正文内容
        letter_number: 函号
        date: 日期（默认今天）
        agent: 经办人
        template_path: 模板文件路径
        output_path: 输出文件路径
    """
    
    # 默认模板路径
    if template_path is None:
        template_path = os.path.join(os.path.dirname(__file__), '..', '请示函_模板.docx')
        if not os.path.exists(template_path):
            template_path = '请示函_模板.docx'
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"模板文件未找到：{template_path}")
    
    # 生成输出路径
    if output_path is None:
        safe_title = "".join([c for c in title if c.isalnum() or c in (' ', '_', '-')]).strip()
        output_path = f'请示_{safe_title}_{letter_number}号.docx'
    
    # 复制模板
    shutil.copy(template_path, output_path)
    
    # 打开文档
    doc = Document(output_path)
    
    # 处理日期
    date_chinese = convert_to_chinese_date(date)
    
    # 构建替换内容
    letter_number_text = f'{letter_number}'
    title_text = f'关于{title}的请示'
    
    # 1. 替换段落中的占位符
    for paragraph in doc.paragraphs:
        # 跳过空段落
        if not paragraph.text.strip():
            continue
        
        # 替换函号
        if replace_placeholder(paragraph, '[函号]', letter_number_text, 15):
            continue
        
        # 替换标题
        if replace_placeholder(paragraph, '[标题]', title_text, 18, font_name='黑体'):
            continue
        
        # 替换正文
        if replace_placeholder(paragraph, '[正文]', content, 15):
            paragraph.paragraph_format.first_line_indent = Pt(28)
            continue
        
        # 替换日期
        if replace_placeholder(paragraph, '[日期]', date_chinese, 15):
            continue
    
    # 2. 替换表格中的占位符（第二页审批表格）
    for table in doc.tables:
        replace_table_placeholder(table, '[函号]', letter_number_text)
        replace_table_placeholder(table, '[标题]', title_text)
        replace_table_placeholder(table, '[经办人]', agent)
        replace_table_placeholder(table, '[日期]', date_chinese)
    
    doc.save(output_path)
    return os.path.abspath(output_path)


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='生成请示函')
    parser.add_argument('--title', required=True, help='事项标题')
    parser.add_argument('--content', required=True, help='正文内容')
    parser.add_argument('--number', required=True, help='函号')
    parser.add_argument('--date', default='今天', help='日期')
    parser.add_argument('--agent', default='游洋洋', help='经办人')
    parser.add_argument('--template', help='模板文件路径')
    parser.add_argument('--output', help='输出文件路径')
    parser.add_argument('--pdf', action='store_true', help='同时生成 PDF 文件')
    
    args = parser.parse_args()
    
    try:
        result = generate_request_letter(
            title=args.title,
            content=args.content,
            letter_number=args.number,
            date=args.date,
            agent=args.agent,
            template_path=args.template,
            output_path=args.output
        )
        print(f"文档已生成：{result}")
        
        # 如果指定了 -pdf 参数，转换为 PDF
        if args.pdf:
            pdf_path = convert_to_pdf(result)
            print(f"PDF已生成：{pdf_path}")
    except Exception as e:
        print(f"错误：{e}", file=sys.stderr)
        sys.exit(1)
